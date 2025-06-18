import os
import tempfile
import chainlit as cl
from PyPDF2 import PdfReader
from docx import Document
import pandas as pd
import requests
from bs4 import BeautifulSoup
import logging
import shutil
import json
import xml.etree.ElementTree as ET
import mimetypes
from pathlib import Path
import zipfile
import csv
from datetime import datetime
import re
from urllib.parse import urlparse, urljoin
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.chrome.service import Service
import time
from typing import Dict, List
from openai import OpenAI

logger = logging.getLogger(__name__)

# Initialize OpenAI client
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

async def handle_file_upload(file: cl.File) -> str:
    """Simplified file upload handler that extracts content and uses GPT-4o for analysis."""
    if not file or not file.path:
        logger.error(f"Invalid file or empty path: {file}")
        return "Error: Invalid file or empty path"

    try:
        file_extension = os.path.splitext(file.name)[1].lower()
        file_size = os.path.getsize(file.path)
        
        logger.info(f"Processing file: {file.name} ({file_extension}, {file_size} bytes)")
        
        # Extract content based on file type
        content = ""
        file_type = "unknown"
        
        if file_extension in ['.py', '.js', '.html', '.css', '.txt', '.md', '.yaml', '.yml', '.json', '.xml', '.sql', '.sh', '.bat']:
            content = extract_text_content(file.path)
            file_type = "text/code"
        elif file_extension == '.pdf':
            content = extract_pdf_content(file.path)
            file_type = "PDF document"
        elif file_extension in ['.docx', '.doc']:
            content = extract_word_content(file.path)
            file_type = "Word document"
        elif file_extension in ['.xlsx', '.xls', '.csv']:
            content = extract_spreadsheet_content(file.path, file_extension)
            file_type = "spreadsheet"
        else:
            return f"⚠️ **Unsupported File Format**\n\n**File:** {file.name}\n**Format:** {file_extension}\n\n**Supported Formats:**\n- Text/Code: TXT, PY, JS, HTML, CSS, MD, JSON, XML, SQL\n- Documents: PDF, DOCX, DOC\n- Spreadsheets: XLSX, XLS, CSV"
        
        if not content or len(content.strip()) < 10:
            return f"❌ **No Content Extracted**\n\n**File:** {file.name}\n**Issue:** Could not extract readable content from this file."
        
        # Use GPT-4o to analyze the content
        logger.info(f"Analyzing content with GPT-4o for file: {file.name}")
        analysis = await analyze_content_with_gpt4o(content, file.name, file_type)
        
        return analysis
        
    except Exception as e:
        logger.error(f"Error processing file {file.name}: {str(e)}")
        return f"❌ **Error Processing File**\n\n**File:** {file.name}\n**Error:** {str(e)}\n\nPlease try uploading the file again or contact support if the issue persists."

def extract_text_content(file_path: str) -> str:
    """Extract content from text-based files."""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception as e:
        logger.error(f"Error reading text file: {str(e)}")
        return ""

def extract_pdf_content(file_path: str) -> str:
    """Extract content from PDF files."""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            content = ""
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                if page_text.strip():
                    content += f"--- Page {page_num} ---\n{page_text}\n\n"
            return content
    except Exception as e:
        logger.error(f"Error reading PDF file: {str(e)}")
        return ""

def extract_word_content(file_path: str) -> str:
    """Extract content from Word documents."""
    try:
        doc = Document(file_path)
        content = ""
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                content += para.text + "\n"
        
        # Extract tables if present
        if doc.tables:
            content += "\n--- Tables ---\n"
            for table_num, table in enumerate(doc.tables, 1):
                content += f"\nTable {table_num}:\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        content += row_text + "\n"
        
        return content
    except Exception as e:
        logger.error(f"Error reading Word document: {str(e)}")
        return ""

def extract_spreadsheet_content(file_path: str, file_extension: str) -> str:
    """Extract content from spreadsheet files."""
    try:
        if file_extension == '.csv':
            df = pd.read_csv(file_path)
        else:
            df = pd.read_excel(file_path)
        
        # Get basic info
        rows, cols = df.shape
        content = f"Spreadsheet Summary:\n"
        content += f"- Dimensions: {rows} rows × {cols} columns\n"
        content += f"- Columns: {', '.join(df.columns.tolist())}\n\n"
        
        # Add sample data (first 10 rows)
        content += "Sample Data (first 10 rows):\n"
        content += df.head(10).to_string(index=False)
        
        # Add summary statistics for numeric columns
        numeric_cols = df.select_dtypes(include=['number']).columns
        if len(numeric_cols) > 0:
            content += "\n\nNumeric Column Statistics:\n"
            content += df[numeric_cols].describe().to_string()
        
        return content
    except Exception as e:
        logger.error(f"Error reading spreadsheet: {str(e)}")
        return ""

async def analyze_content_with_gpt4o(content: str, filename: str, file_type: str) -> str:
    """Use GPT-4o to analyze file content and provide intelligent summary."""
    try:
        # Limit content size for API efficiency (keep first 8000 characters)
        content_snippet = content[:8000] if len(content) > 8000 else content
        
        # Create analysis prompt
        prompt = f"""Analyze this {file_type} file named "{filename}" and provide a comprehensive summary.

File Content:
{content_snippet}

Please provide:
1. **Document Summary** - What this file is about and its main purpose
2. **Key Points** - Main topics, findings, or important information (use bullet points)
3. **Technical Details** - Any technical specifications, code functionality, or data insights
4. **Structure & Organization** - How the content is organized
5. **Actionable Insights** - What someone could learn or do with this information

Format your response with clear headings and bullet points for easy reading. Be concise but comprehensive."""
        
        # Make API call to GPT-4o
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "system", 
                    "content": "You are an expert document analyzer. Provide clear, structured summaries that help users quickly understand the content and value of their documents. Use bullet points, clear headings, and highlight key information."
                },
                {
                    "role": "user", 
                    "content": prompt
                }
            ],
            temperature=0.7,
            max_tokens=1500
        )
        
        analysis = response.choices[0].message.content
        
        # Add file metadata
        file_info = f"""# 📄 **File Analysis: {filename}**

**📋 File Information:**
• **Type:** {file_type}
• **Size:** {len(content):,} characters
• **Analysis:** GPT-4o powered analysis

---

{analysis}

---

**💡 Need more details?** Ask me specific questions about this file content!"""
        
        return file_info
        
    except Exception as e:
        logger.error(f"Error in GPT-4o analysis: {str(e)}")
        
        # Fallback to basic summary
        word_count = len(content.split())
        char_count = len(content)
        
        return f"""# 📄 **File Analysis: {filename}**

**📋 File Information:**
• **Type:** {file_type}
• **Size:** {char_count:,} characters ({word_count:,} words)
• **Status:** ⚠️ AI analysis unavailable, showing basic summary

**📝 Content Preview:**
{content[:500]}{'...' if len(content) > 500 else ''}

**❌ Analysis Error:** {str(e)}

Please try again or ask specific questions about the file content."""

async def handle_url(url: str) -> str:
    """Simplified URL content extraction."""
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Remove unwanted elements
        for element in soup(["script", "style", "nav", "header", "footer"]):
            element.decompose()
        
        # Extract title and content
        title = soup.find('title')
        title_text = title.get_text().strip() if title else "No title"
        
        content = soup.get_text(separator=' ', strip=True)
        
        # Use GPT-4o to analyze the URL content
        analysis = await analyze_content_with_gpt4o(content, url, "webpage")
        
        return f"🌐 **URL Analysis: {url}**\n\n**Title:** {title_text}\n\n{analysis}"
        
    except Exception as e:
        logger.error(f"Error processing URL {url}: {str(e)}")
        return f"❌ **Error Processing URL**\n\n**URL:** {url}\n**Error:** {str(e)}"
